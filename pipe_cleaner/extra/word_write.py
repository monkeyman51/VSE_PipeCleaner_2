from shutil import copyfile
from docx import Document
from datetime import datetime


def copy_word_document(pipe_num: str, full_name: str):
    new_document = Document()
    word_file_name = f'task_{pipe_num}.docx'
    word_doc_path = f'pipes/{full_name}/{word_file_name}'

    new_document.save(word_doc_path)

    copyfile('input/technician_lab_template.docx', word_doc_path)

    return word_doc_path


def edit_word_document(user_name: str, pipe_num: str, location: str, start_date: str, due_date: str,
                       word_file_path: str):
    tech_document = Document(word_file_path)
    date_today = datetime.today().strftime('%Y-%m-%d')
    time_now = datetime.today().strftime('%H:%M')

    dict_1 = {'user_name': user_name,
              'pipe_number': pipe_num,
              'pipe_location': location,
              'start_date': 'Some Date',
              'due_date': 'Some Date',
              'date_today': date_today,
              'time_now': time_now}

    dict_2 = {' user_name': f' {user_name}',
              ' pipe_number': f' {pipe_num}',
              ' pipe_location': f' {location}',
              ' start_date': ' 10/14/2020 7:00 AM',
              ' due_date': ' 11/6/2020 8:00 AM',
              ' date_today': date_today,
              ' time_now': time_now}
    # for paragraph in tech_document.paragraphs:
    #     inline = paragraph.runs
    #     for i in range(len(inline)):
    #         text = inline[i].text
    #         if text in dict_1.keys():
    #             text = text.replace(text, dict_1[text])
    #             inline[i].text = text
    #         elif text in dict_2.keys():
    #             text = text.replace(text, dict_2[text])
    #             inline[i].text = text


    section = tech_document.sections[0]
    header = section.header
    for paragraph in header.paragraphs:
        inline = paragraph.runs
        for i in range(len(inline)):
            text = inline[i].text
            if text in dict_1.keys():
                text = text.replace(text, dict_1[text])
                inline[i].text = text
            # elif 'date_today' in text:
            #     text = text.replace(text, dict_1[text])
            #     inline[i].text = text
            elif text in dict_2.keys():
                text = text.replace(text, dict_2[text])
                inline[i].text = text

    tech_document.save(word_file_path)


def main_method(user_name, pipe_name: str, full_name: str):
    word_doc_path = copy_word_document(pipe_name, full_name)
    edit_word_document(user_name, pipe_name, 'R21', 'Some', 'Some', word_doc_path)

    return word_doc_path
